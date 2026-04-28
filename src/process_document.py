# src/process_document.py

import os
import json
import asyncio
from pathlib import Path
from src.ai_search import (
    index_document,
    build_text_for_ai,
    build_chunk_context_prefix,
    add_document_page_chunks,
    add_document_chunks,
    save_chunk_index
)
from .db_utils import save_document_async
from datetime import datetime
from src.main import parse_document_with_gpt
from src.document_field_flattener import flatten_document_fields
from src.document_field_normalizer import (
    extract_canonical_fields_only,
    normalize_document_type,
)    
UPLOADS_DIR = Path(__file__).parent / "uploads"

def extract_text_universal(file_path: str):
    import os
    from docx import Document
    from .main import extract_text, extract_text_pages_from_pdf

    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        page_texts = extract_text_pages_from_pdf(file_path)
        page_texts = [p for p in (page_texts or []) if p and p.strip()]

        if page_texts:
            return "\n\n".join(page_texts), page_texts

        text = extract_text(file_path)
        return text, [text] if text else []

    if ext == ".docx":
        doc = Document(file_path)
        parts = []

        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text.strip())

        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

        text = "\n".join(parts)
        return text, [text] if text else []

    if ext in [".jpg", ".jpeg", ".png", ".bmp", ".tiff", ".webp"]:
        text = extract_text(file_path)
        return text, [text] if text else []

    text = extract_text(file_path)
    return text, [text] if text else []
    
def estrai_testo_da_file(file_path: str) -> str:
    """
    Estrae solo il testo da un file PDF o immagine.
    Restituisce il testo come stringa.
    """
    ext = os.path.splitext(file_path)[1].lower()

    from .main import extract_text, extract_text_from_pdf

    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    else:
        return extract_text(file_path)
        
async def process_document(file_path: str):
    from .main import (
        extract_text,
        extract_text_pages_from_pdf,
        parse_document_with_gpt,
        generate_word
    )

    print(f"[PROCESS] Elaborazione documento: {file_path}")

    ext = os.path.splitext(file_path)[1].lower()

    text, page_texts = await asyncio.to_thread(
        extract_text_universal,
        file_path
    )

    is_passport = False
    mrz_data = {}

    if "passport" in text.lower() or "passeport" in text.lower():
        is_passport = True
        print("[PROCESS] Documento rilevato come PASSAPORTO")

    if is_passport and ext in [".jpg", ".jpeg", ".png", ".bmp", ".tiff", ".webp"]:
        from .main import extract_mrz_from_image, parse_mrz

        mrz_text = await asyncio.to_thread(extract_mrz_from_image, file_path)
        mrz_data = parse_mrz(mrz_text)
        print("[MRZ PARSED]:", mrz_data)

    debug_dir = Path("debug")
    debug_dir.mkdir(exist_ok=True)

    ts = datetime.now().strftime("%Y%m%d_%H%M%S_%f")
    debug_path = debug_dir / f"debug_ocr_{Path(file_path).stem}_{ts}.txt"

    with open(debug_path, "w", encoding="utf-8") as f:
        f.write(text)

    # 1. Parsing unificato: tipo documento + campi completi
    parsed_doc_type, parsed_fields = await asyncio.to_thread(
        parse_document_with_gpt,
        text
    )

    doc_type = normalize_document_type(parsed_doc_type)

    print(f"[PROCESS] Tipo documento unificato: {doc_type}")
    print(f"[PROCESS] Unified fields: {parsed_fields}")

    # 2. Fonte unica dei dati
    fields = dict(parsed_fields or {})

    # 5. Merge MRZ se presente
    if mrz_data:
        print("[PROCESS] Merge MRZ nei campi")
        print("[PROCESS] MRZ data:", mrz_data)
        fields.update(mrz_data)

    # 6. Salva il nome completo originale prima del flatten
    original_full_name = (
        fields.get("nome_completo_originale")
        or fields.get("nome_completo")
        or fields.get("full_name")
        or fields.get("cognome_nome")
        or fields.get("nome")
    )

    print(f"[PROCESS] Campi grezzi finali: {fields}")

    # 7. Flatten universale del payload ricco
    flat_fields = flatten_document_fields(fields)
    print(f"[PROCESS] Campi ricchi flattenati: {flat_fields}")

    # 8. Campi canonici SOLO per search/business rules
    canonical_fields = extract_canonical_fields_only(doc_type, flat_fields)
    print(f"[PROCESS] Campi canonici (solo search/business): {canonical_fields}")

    if (
        canonical_fields.get("ore_lavorate_totale") is not None
        and canonical_fields.get("ore_lavorate_lavoro_ordinario") is not None
        and str(canonical_fields["ore_lavorate_totale"]).strip() == str(canonical_fields["ore_lavorate_lavoro_ordinario"]).strip()
    ):
        canonical_fields.pop("ore_lavorate_totale", None)

    # 9. Campi finali da salvare/mostrare = SOLO payload ricco flattenato
    final_fields = dict(flat_fields or {})
    print(f"[PROCESS] Campi salvati nel DB (payload ricco): {final_fields}")

    if original_full_name and "nome_completo_originale" not in final_fields:
        final_fields["nome_completo_originale"] = original_full_name
    
    
    if (
        final_fields.get("ore_lavorate_totale") is not None
        and final_fields.get("ore_lavorate_lavoro_ordinario") is not None
        and str(final_fields["ore_lavorate_totale"]).strip() == str(final_fields["ore_lavorate_lavoro_ordinario"]).strip()
    ):
        final_fields.pop("ore_lavorate_totale", None)

    print(f"[PROCESS] Campi finali persistiti: {final_fields}")

    ocr_pages_json = json.dumps(page_texts, ensure_ascii=False)

    doc_id = await save_document_async(
        doc_type,
        file_path,
        final_fields,
        canonical_fields=canonical_fields,
        ocr_text=text,
        ocr_pages=ocr_pages_json
    )
    print(f"[PROCESS] Documento salvato con ID: {doc_id}")

    text_for_ai = build_text_for_ai(
        doc_type,
        file_path,
        final_fields,
        canonical_fields=canonical_fields,
        ocr_text=text
    )

    chunk_context_prefix = build_chunk_context_prefix(
        doc_type,
        file_path,
        final_fields,
        canonical_fields=canonical_fields
    )

    await asyncio.to_thread(index_document, doc_id, text_for_ai)

    if page_texts:
        await asyncio.to_thread(add_document_page_chunks, doc_id, page_texts, chunk_context_prefix)
    else:
        await asyncio.to_thread(add_document_chunks, doc_id, text, chunk_context_prefix)

    await asyncio.to_thread(save_chunk_index)
    print(f"[PROCESS] Chunk index salvato per doc_id={doc_id}")

    word_path = await asyncio.to_thread(generate_word, final_fields, file_path)
    print(f"[PROCESS] File Word generato: {word_path}")

    return doc_id, word_path