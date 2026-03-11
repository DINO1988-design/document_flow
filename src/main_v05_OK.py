import os
import time
import shutil
from datetime import datetime
from watchdog.observers import Observer
from watchdog.events import FileSystemEventHandler
from PIL import Image
import pytesseract
from docx import Document
from pdf2image import convert_from_path
import openai
import json

# --------------------------
# CONFIGURAZIONE API GPT
# --------------------------
openai.api_key = os.getenv("OPENAI_API_KEY")

# --------------------------
# PERCORSO TESSERACT
# --------------------------
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

# --------------------------
# DATABASE
# --------------------------
from db_setup import engine, documents, document_fields

# --------------------------
# CARTELLE
# --------------------------
INPUT_FOLDER = "input/"
PROCESSED_FOLDER = "processed/"
OUTPUT_FOLDER = "output/"

# ==========================
# FUNZIONI DATABASE
# ==========================
def save_document(tipo_documento, file_path, fields):

    try:
        with engine.begin() as conn:

            result = conn.execute(
                documents.insert().values(
                    tipo_documento=tipo_documento,
                    file_path=file_path,
                    data_creazione=datetime.now()
                )
            )

            document_id = result.inserted_primary_key[0]

            for campo, valore in fields.items():

                if isinstance(valore, (dict, list)):
                    valore = json.dumps(valore, ensure_ascii=False)

                if isinstance(valore, str) and len(valore) > 4000:
                    valore = valore[:4000]

                conn.execute(
                    document_fields.insert().values(
                        document_id=document_id,
                        campo=campo,
                        valore=valore
                    )
                )

        return document_id

    except Exception as e:
        print("Errore salvataggio DB:", e)
        return None


# ==========================
# OCR
# ==========================

def preprocess_image(img):
    """Migliora immagine per OCR"""
    img = img.convert("L")
    return img


def extract_text(file_path):

    try:
        img = Image.open(file_path)
        img = preprocess_image(img)

        text = pytesseract.image_to_string(
            img,
            lang="ita+eng"
        )

        return text

    except Exception as e:
        print(f"Errore OCR immagine {file_path}: {e}")
        return ""


def extract_text_from_pdf(pdf_path):

    try:

        text_all = ""

        pages = convert_from_path(
            pdf_path,
            poppler_path=r"C:\Program Files\poppler-25.12.0\bin"
        )

        for page in pages:

            page = preprocess_image(page)

            text_page = pytesseract.image_to_string(
                page,
                lang="ita+eng"
            )

            text_all += text_page + "\n"

        return text_all

    except Exception as e:
        print(f"Errore OCR PDF {pdf_path}: {e}")
        return ""


# ==========================
# PULIZIA TESTO OCR
# ==========================

def clean_ocr_text(text):

    text = text.replace("\n\n", "\n")
    text = text.replace("  ", " ")
    text = text.strip()

    max_chars = 12000

    if len(text) > max_chars:
        text = text[:max_chars]

    return text


# ==========================
# GPT PARSING
# ==========================

def parse_data_with_gpt(text):

    text = clean_ocr_text(text)

    system_prompt = """
Sei un sistema professionale di document understanding.

Riceverai il testo OCR di un documento.

Devi:

1) Identificare il tipo di documento
esempi:
- fattura
- ricevuta
- contratto
- carta_identita
- passaporto
- referto_medico
- certificato
- preventivo
- modulo
- documento_generico

2) Estrarre tutti i dati rilevanti.

Linee guida:

- usa chiavi descrittive in italiano
- non inventare dati
- se non sei sicuro non inserire il campo
- estrai numeri, date, importi se presenti

Rispondi SOLO con JSON valido.

Formato:

{
 "tipo_documento": "stringa",
 "campi": {
    "nome_campo": "valore"
 }
}
"""

    user_prompt = f"""
Testo OCR del documento:

{text}
"""

    for attempt in range(2):

        try:

            response = openai.chat.completions.create(
                model="gpt-4o-mini",
                temperature=0,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ]
            )

            content = response.choices[0].message.content.strip()

            if content.startswith("```"):
                content = content.replace("```json", "").replace("```", "").strip()

            data = json.loads(content)

            tipo_documento = data.get("tipo_documento", "generico")
            fields = data.get("campi", {})

            if not isinstance(fields, dict):
                fields = {}

            return tipo_documento, fields

        except Exception as e:
            print(f"Errore GPT tentativo {attempt+1}: {e}")

    return "generico", {}


# ==========================
# GENERAZIONE WORD
# ==========================

def generate_word(fields, original_file):

    try:

        doc = Document()
        doc.add_heading("Dati Documento", 0)

        for campo, valore in fields.items():

            if isinstance(valore, (dict, list)):
                valore = json.dumps(valore, ensure_ascii=False)

            doc.add_paragraph(f"{campo}: {valore}")

        output_path = os.path.join(
            OUTPUT_FOLDER,
            os.path.basename(original_file) + ".docx"
        )

        doc.save(output_path)

        print("Word generato:", output_path)

    except Exception as e:
        print("Errore generazione Word:", e)


# ==========================
# PROCESSO DOCUMENTO
# ==========================

def process_document(file_path):

    time.sleep(1)

    print("\n--------------------------")
    print("Nuovo documento:", file_path)

    text = ""

    if file_path.lower().endswith((".jpg", ".jpeg", ".png")):

        text = extract_text(file_path)

    elif file_path.lower().endswith(".pdf"):

        text = extract_text_from_pdf(file_path)

    else:

        print("Formato non supportato")
        return

    if not text.strip():

        print("OCR non ha trovato testo")

    tipo_documento, fields = parse_data_with_gpt(text)

    print("Tipo documento:", tipo_documento)
    print("Campi trovati:", len(fields))

    save_document(tipo_documento, file_path, fields)

    generate_word(fields, file_path)

    shutil.move(
        file_path,
        os.path.join(PROCESSED_FOLDER, os.path.basename(file_path))
    )

    print("File spostato in processed")


# ==========================
# WATCHER CARTELLA
# ==========================

class Watcher(FileSystemEventHandler):

    def on_created(self, event):

        if not event.is_directory:

            process_document(event.src_path)


# ==========================
# MAIN
# ==========================

if __name__ == "__main__":

    os.makedirs(INPUT_FOLDER, exist_ok=True)
    os.makedirs(PROCESSED_FOLDER, exist_ok=True)
    os.makedirs(OUTPUT_FOLDER, exist_ok=True)

    observer = Observer()

    observer.schedule(
        Watcher(),
        INPUT_FOLDER,
        recursive=False
    )

    observer.start()

    print("\nWatcher avviato")
    print("Metti file in /input")

    try:

        while True:
            time.sleep(1)

    except KeyboardInterrupt:

        observer.stop()

    observer.join()