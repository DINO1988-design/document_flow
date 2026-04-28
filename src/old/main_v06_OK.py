from sqlalchemy.orm import Session
import os
import json
import shutil
from datetime import datetime
from PIL import Image
import pytesseract
from pdf2image import convert_from_path
from docx import Document
import openai
from src.db_setup import engine, documents, document_fields

# --------------------------
# Configurazione
# --------------------------
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
UPLOAD_FOLDER = "src/uploads/"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

openai.api_key = os.getenv("OPENAI_API_KEY")

# --------------------------
# Funzioni OCR
# --------------------------
def extract_text(file_path):
    try:
        img = Image.open(file_path).convert("L")
        return pytesseract.image_to_string(img, lang="ita+eng")
    except Exception as e:
        print(f"Errore OCR {file_path}: {e}")
        return ""

def extract_text_from_pdf(pdf_path):
    try:
        text_all = ""
        pages = convert_from_path(pdf_path, poppler_path=r"C:\Program Files\poppler-25.12.0\bin")
        for page in pages:
            text_all += pytesseract.image_to_string(page, lang="ita+eng") + "\n"
        return text_all
    except Exception as e:
        print(f"Errore OCR PDF {pdf_path}: {e}")
        return ""

# --------------------------
# Funzioni GPT
# --------------------------
def clean_ocr_text(text):
    text = text.replace("\n\n", "\n").replace("  ", " ").strip()
    return text[:12000]  # limite sicurezza token

def parse_data_with_gpt(text):
    text = clean_ocr_text(text)
    system_prompt = """
Sei un sistema professionale di document understanding.
Devi determinare il tipo di documento e restituire TUTTI i dati rilevanti.
Rispondi SOLO con JSON:
{
 "tipo_documento": "stringa",
 "campi": {}
}
"""
    user_prompt = f"Testo OCR:\n{text}"
    try:
        response = openai.chat.completions.create(
            model="gpt-4o-mini",
            temperature=0,
            messages=[{"role":"system","content":system_prompt},{"role":"user","content":user_prompt}]
        )
        content = response.choices[0].message.content.strip()
        if content.startswith("```"):
            content = content.replace("```json","").replace("```","").strip()
        data = json.loads(content)
        tipo_documento = data.get("tipo_documento","generico")
        fields = data.get("campi",{})
        if not isinstance(fields, dict):
            fields = {}
        return tipo_documento, fields
    except Exception as e:
        print(f"Errore GPT parsing: {e}")
        return "generico", {}

# --------------------------
# Salvataggio database
# --------------------------
def save_document(tipo_documento, file_path, fields):
    try:
        with Session(engine) as session:
            result = session.execute(
                documents.insert().values(
                    tipo_documento=tipo_documento,
                    file_path=file_path,
                    data_creazione=datetime.now()
                )
            )
            document_id = result.inserted_primary_key[0]

            for campo, valore in fields.items():
                if isinstance(valore, (dict, list)):
                    valore = json.dumps(valore, ensure_ascii=False)
                if isinstance(valore, str) and len(valore) > 4000:
                    valore = valore[:4000]

                session.execute(
                    document_fields.insert().values(
                        document_id=document_id,
                        campo=campo,
                        valore=valore
                    )
                )

            session.commit()
            print(f"Documento salvato con ID {document_id}")
        return document_id
    except Exception as e:
        print("Errore salvataggio DB:", e)
        return None

# --------------------------
# Generazione Word
# --------------------------
def generate_word(fields, original_file):
    os.makedirs("output", exist_ok=True)
    doc = Document()
    doc.add_heading("Dati Documento", 0)
    for campo, valore in fields.items():
        if isinstance(valore,(dict,list)):
            valore = json.dumps(valore, ensure_ascii=False)
        doc.add_paragraph(f"{campo}: {valore}")
    output_path = os.path.join("output", os.path.basename(original_file)+".docx")
    doc.save(output_path)
    print(f"Documento Word generato: {output_path}")

# --------------------------
# Processo completo
# --------------------------
def process_document(file_path):
    print(f"Elaborazione {file_path}")
    text = ""
    if file_path.lower().endswith((".jpg",".jpeg",".png")):
        text = extract_text(file_path)
    elif file_path.lower().endswith(".pdf"):
        text = extract_text_from_pdf(file_path)
    else:
        print("Formato non supportato")
        return
    tipo_documento, fields = parse_data_with_gpt(text)
    save_document(tipo_documento,file_path,fields)
    generate_word(fields,file_path)