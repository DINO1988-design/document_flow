# main.py
from sqlalchemy.orm import Session
import os
import json
from datetime import datetime
from PIL import Image
import pytesseract
import re
from pdf2image import convert_from_path
from docx import Document
import openai
import pdfplumber
from src.db_setup import async_engine, documents, document_fields
from sqlalchemy import insert
from src.db_utils import save_document_async
from src.document_field_normalizer import (
    get_document_schema,
    get_allowed_document_types,
    get_allowed_document_types_for_prompt,
)
   
# --------------------------
# Configurazione OCR
# --------------------------
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
UPLOAD_FOLDER = "src/uploads/"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

openai.api_key = os.getenv("OPENAI_API_KEY")

# --------------------------
# Funzioni OCR
# --------------------------
def extract_text(file_path):
    try:
        img = Image.open(file_path).convert("L")
        return pytesseract.image_to_string(img, lang="ita+eng")
    except Exception as e:
        print(f"Errore OCR {file_path}: {e}")
        return ""

def extract_mrz_from_image(file_path):
    try:
        img = Image.open(file_path)

        width, height = img.size

        # 👉 prende solo la parte bassa (MRZ)
        mrz_crop = img.crop((0, int(height * 0.75), width, height))

        mrz_text = pytesseract.image_to_string(
            mrz_crop,
            lang="eng",
            config="--psm 6"
        )

        print("\n========== MRZ OCR ==========\n")
        print(mrz_text)
        print("\n=============================\n")

        return mrz_text

    except Exception as e:
        print(f"Errore MRZ OCR: {e}")
        return ""

def parse_mrz(mrz_text):
    try:
        lines = [l.strip() for l in mrz_text.split("\n") if "<" in l]

        if len(lines) < 2:
            return {}

        line1 = lines[0]
        line2 = lines[1]

        # Nome e cognome
        names = line1.split("<<")
        surname = names[0][5:].replace("<", " ").strip()
        given_names = names[1].replace("<", " ").strip() if len(names) > 1 else ""

        # Numero passaporto
        passport_number = line2[0:9].replace("<", "")

        # Nazionalità
        nationality = line2[10:13]

        # Data nascita
        birth_date = line2[13:19]

        # Sesso
        sex = line2[20]

        # Data scadenza
        expiry = line2[21:27]

        return {
            "cognome": surname,
            "nome": given_names,
            "numero_passaporto": passport_number,
            "nazionalita": nationality,
            "data_nascita": birth_date,
            "sesso": sex,
            "data_scadenza": expiry
        }

    except Exception as e:
        print("Errore parsing MRZ:", e)
        return {}
        
def extract_pdf_pages_native(pdf_path):
    pages_text = []
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text() or ""
                pages_text.append(text)
    except Exception as e:
        print(f"Errore estrazione testo nativo PDF {pdf_path}: {e}")
        return []
    return pages_text


def extract_pdf_pages_ocr(pdf_path):
    pages_text = []
    try:
        pages = convert_from_path(
            pdf_path,
            poppler_path=r"C:\Program Files\poppler-25.12.0\bin"
        )
        for page in pages:
            text = pytesseract.image_to_string(page, lang="ita+eng")
            pages_text.append(text)
    except Exception as e:
        print(f"Errore OCR PDF {pdf_path}: {e}")
        return []
    return pages_text




def is_good_text(text: str) -> bool:
    if not text:
        return False

    # Pulizia base
    clean = re.sub(r"\s+", " ", text).strip()

    # 1. Lunghezza minima
    if len(clean) < 120:
        return False

    words = clean.split()

    # 2. Numero minimo parole
    if len(words) < 25:
        return False

    # 3. Conta numeri (dati reali)
    digit_count = sum(ch.isdigit() for ch in clean)

    # 4. Parole tutte maiuscole (tipico layout PDF)
    upper_words = sum(1 for w in words if len(w) > 2 and w.isupper())
    upper_ratio = upper_words / max(len(words), 1)

    # 5. Parole “naturali” (tipo Mario, Rossi)
    mixed_words = sum(1 for w in words if re.search(r"[A-Z]", w) and re.search(r"[a-z]", w))
    mixed_ratio = mixed_words / max(len(words), 1)

    # 6. Pattern di dati veri
    signal_patterns = 0

    # date
    if re.search(r"\b\d{2}[/-]\d{2}[/-]\d{2,4}\b", clean):
        signal_patterns += 1

    # importi
    if re.search(r"\b\d+[.,]\d{2}\b", clean):
        signal_patterns += 1

    # email
    if "@" in clean:
        signal_patterns += 1

    # IBAN (approssimato)
    if re.search(r"\b[A-Z]{2}\d{2}[A-Z0-9]{10,30}\b", clean, re.I):
        signal_patterns += 1

    # codice fiscale (approssimato)
    if re.search(r"\b[A-Z]{6}\d{2}[A-Z]\d{2}[A-Z]\d{3}[A-Z]\b", clean, re.I):
        signal_patterns += 1

    # ---------------------------
    # REGOLE DECISIONALI
    # ---------------------------

    # ❌ troppo layout, pochi dati
    if upper_ratio > 0.45 and digit_count < 20 and signal_patterns == 0:
        return False

    # ❌ testo quasi tutto intestazioni
    if upper_ratio > 0.55 and mixed_ratio < 0.05:
        return False

    # ❌ pochissimi segnali reali
    if digit_count < 10 and signal_patterns == 0:
        return False

    return True


def extract_text_from_pdf(pdf_path):
    native_pages = extract_pdf_pages_native(pdf_path)

    if not native_pages:
        print("[PDF] Nessun testo nativo trovato, uso OCR completo")
        return "\n\n".join(extract_pdf_pages_ocr(pdf_path))

    final_pages = []
    ocr_pages = None

    for i, native_text in enumerate(native_pages):
        if is_good_text(native_text):
            print(f"[PDF] Pagina {i+1}: testo nativo OK")
            final_pages.append(native_text)
        else:
            print(f"[PDF] Pagina {i+1}: testo nativo debole, uso OCR")
            if ocr_pages is None:
                ocr_pages = extract_pdf_pages_ocr(pdf_path)

            ocr_text = ocr_pages[i] if i < len(ocr_pages) else ""
            final_pages.append(ocr_text)

    return "\n\n".join(final_pages)

def extract_text_pages_from_pdf(pdf_path):
    native_pages = extract_pdf_pages_native(pdf_path)

    if not native_pages:
        print("[PDF] Nessun testo nativo trovato, uso OCR completo per pagine")
        return extract_pdf_pages_ocr(pdf_path)

    final_pages = []
    ocr_pages = None

    for i, native_text in enumerate(native_pages):
        if is_good_text(native_text):
            print(f"[PDF] Pagina {i+1}: testo nativo OK")
            final_pages.append(native_text)
        else:
            print(f"[PDF] Pagina {i+1}: testo nativo debole, uso OCR")
            if ocr_pages is None:
                ocr_pages = extract_pdf_pages_ocr(pdf_path)

            ocr_text = ocr_pages[i] if i < len(ocr_pages) else ""
            final_pages.append(ocr_text)

    return final_pages
    
# --------------------------
# Funzioni GPT + max lenghth tocken
# --------------------------
def clean_ocr_text(text: str, max_len: int = 80000) -> str:
    if not text:
        return ""

    # 1. Normalizza ritorni a capo
    text = text.replace("\r\n", "\n").replace("\r", "\n")

    # 2. Rimuove caratteri di controllo non utili
    text = "".join(
        ch for ch in text
        if ch == "\n" or ch == "\t" or ord(ch) >= 32
    )

    # 3. Normalizza spazi e tab
    text = text.replace("\t", " ")
    text = re.sub(r"[ ]{2,}", " ", text)

    # 4. Riduce righe vuote multiple
    text = re.sub(r"\n{3,}", "\n\n", text)

    # 5. Pulisce spazi ai bordi delle righe
    lines = [line.strip() for line in text.split("\n")]

    cleaned_lines = []
    for line in lines:
        if not line:
            cleaned_lines.append("")
            continue

        # 6. Riduce rumore OCR molto comune senza toccare dati utili
        line = line.replace("’", "'").replace("“", '"').replace("”", '"')
        line = re.sub(r"[ ]{2,}", " ", line)

        # 7. Separa lettere e numeri attaccati quando chiaramente fusi
        # es: Mario123 -> Mario 123, 123Mario -> 123 Mario
        line = re.sub(r"([A-Za-zÀ-ÖØ-öø-ÿ])(\d)", r"\1 \2", line)
        line = re.sub(r"(\d)([A-Za-zÀ-ÖØ-öø-ÿ])", r"\1 \2", line)

        # 8. Compatta punteggiatura con spazi anomali
        line = re.sub(r"\s+([,.:;])", r"\1", line)
        line = re.sub(r"([,.:;])([^\s])", r"\1 \2", line)

        # 9. Mantieni slash e trattini utili per date/codici, ma togli spazi strani attorno
        line = re.sub(r"\s*/\s*", "/", line)
        line = re.sub(r"\s*-\s*", "-", line)

        cleaned_lines.append(line)

    text = "\n".join(cleaned_lines)

    # 10. Elimina righe quasi vuote o solo simboli ripetuti
    final_lines = []
    for line in text.split("\n"):
        stripped = line.strip()

        if not stripped:
            final_lines.append("")
            continue

        # righe tipo ----====****
        if re.fullmatch(r"[-=*_~.]+", stripped):
            continue

        final_lines.append(stripped)

    text = "\n".join(final_lines)

    # 11. Ricompatta righe vuote multiple dopo la pulizia
    text = re.sub(r"\n{3,}", "\n\n", text).strip()

    # 12. Limite sicurezza token
    return text[:max_len]

def parse_document_with_gpt(text: str):
    original_len = len(text or "")
    text = clean_ocr_text(text or "")

    print(f"[GPT][UNIFIED] Lunghezza OCR originale: {original_len}")
    print(f"[GPT][UNIFIED] Lunghezza OCR inviata a GPT: {len(text)}")

    allowed_types_text = get_allowed_document_types_for_prompt()

    system_prompt = f"""
Sei un sistema professionale di document understanding.

Obiettivi:
1. Identificare il tipo reale del documento.
2. Estrarre il maggior numero possibile di dati realmente presenti nel documento.
3. Restituire una struttura JSON ricca, completa e fedele al contenuto.

Tipi documento ammessi:
{allowed_types_text}

Regole:
- Non inventare dati.
- Estrai tutti i dati utili realmente presenti.
- Puoi usare strutture annidate se aiutano a rappresentare meglio il contenuto.
- Non limitarti ai soli campi principali.
- Se un dato non è presente, non inserirlo.
- Restituisci SEMPRE e SOLO JSON valido.

Formato obbligatorio:
{{
  "tipo_documento": "uno dei tipi ammessi",
  "campi": {{}}
}}
""".strip()

    user_prompt = f"Testo OCR:\n{text}"

    print("\n================ PROMPT GPT UNIFICATO =================")
    print("[SYSTEM PROMPT]")
    print(system_prompt)
    print("\n[USER PROMPT]")
    print(user_prompt)
    print("======================================================\n")

    try:
        response = openai.chat.completions.create(
            model="gpt-4o-mini",
            temperature=0,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ]
        )

        content = response.choices[0].message.content.strip()
        print("[GPT][UNIFIED] Risposta raw:", content)

        if content.startswith("```"):
            content = content.replace("```json", "").replace("```", "").strip()

        data = json.loads(content)

        if not isinstance(data, dict):
            return "generico", {}

        tipo_documento = data.get("tipo_documento", "generico")
        fields = data.get("campi", {}) or {}

        if not isinstance(fields, dict):
            fields = {}

        print("[GPT][UNIFIED] Tipo documento:", tipo_documento)
        print("[GPT][UNIFIED] Payload completo estratto:", fields)

        return tipo_documento, fields

    except Exception as e:
        print(f"Errore GPT unified parsing: {e}")
        return "generico", {}
        
def parse_data_with_gpt_bootstrap(text: str):
    original_len = len(text or "")
    text = clean_ocr_text(text or "")

    print(f"[GPT][BOOTSTRAP] Lunghezza OCR originale: {original_len}")
    print(f"[GPT][BOOTSTRAP] Lunghezza OCR inviata a GPT: {len(text)}")

    system_prompt = """
Sei un sistema professionale di document understanding.

Obiettivo:
1. Capire il tipo reale del documento.
2. Estrarre i dati principali presenti nel documento.
3. Restituire i dati in formato JSON pulito.

Tipi documento possibili:
- cedolino
- bonifico
- biglietto
- passaporto
- fattura
- contratto
- generico

Regole:
- Estrai i dati realmente presenti nel documento.
- Usa nomi di campi semplici, chiari e coerenti.
- Non inventare dati.
- Se un dato non è presente, non inserirlo.
- Restituisci SEMPRE e SOLO JSON valido.

Output obbligatorio (SOLO JSON):
{
  "tipo_documento": "stringa",
  "campi": {}
}
""".strip()

    user_prompt = f"Testo OCR:\n{text}"

    try:
        response = openai.chat.completions.create(
            model="gpt-4o-mini",
            temperature=0,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ]
        )

        content = response.choices[0].message.content.strip()
        print("[GPT][BOOTSTRAP] Risposta raw:", content)

        if content.startswith("```"):
            content = content.replace("```json", "").replace("```", "").strip()

        data = json.loads(content)

        tipo_documento = data.get("tipo_documento", "generico")
        fields = data.get("campi", {})

        if not isinstance(fields, dict):
            fields = {}

        return tipo_documento, fields

    except Exception as e:
        print(f"Errore GPT bootstrap parsing: {e}")
        return "generico", {}
def build_schema_driven_parsing_prompt(document_type: str) -> str:
    schema = get_document_schema(document_type)

    canonical_fields = schema.get("canonical_fields", [])
    required_fields = schema.get("required_fields", [])

    return f"""
Sei un sistema di estrazione dati.

Restituisci JSON.
""".strip()


def parse_data_with_gpt(text: str, document_type: str):
    original_len = len(text or "")
    text = clean_ocr_text(text or "")

    print(f"[GPT][SCHEMA] Lunghezza OCR originale: {original_len}")
    print(f"[GPT][SCHEMA] Lunghezza OCR inviata a GPT: {len(text)}")

    system_prompt = build_schema_driven_parsing_prompt(document_type)
    user_prompt = f"Testo OCR:\n{text}"

    print("\n================ PROMPT GPT =================")
    print("[SYSTEM PROMPT]")
    print(system_prompt)

    print("\n[USER PROMPT]")
    print(user_prompt)

    print("============================================\n")

    try:
        response = openai.chat.completions.create(
            model="gpt-4o-mini",
            temperature=0,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ]
        )

        content = response.choices[0].message.content.strip()
        print("[GPT][SCHEMA] Risposta raw:", content)

        if content.startswith("```"):
            content = content.replace("```json", "").replace("```", "").strip()

        data = json.loads(content)

        if not isinstance(data, dict):
            return document_type, {}

        tipo_documento = data.get("tipo_documento", document_type)

        campi_canonici = data.get("campi_canonici", {}) or {}
        campi_extra = data.get("campi_extra", {}) or {}

        if not isinstance(campi_canonici, dict):
            campi_canonici = {}

        if not isinstance(campi_extra, dict):
            campi_extra = {}

        # ---------------------------------------------------
        # FALLBACK LOSSLESS:
        # se GPT non usa il formato atteso, conserva comunque
        # tutto il payload utile senza perderlo
        # ---------------------------------------------------
        payload_completo = {}

        # caso ideale: campi_extra presente
        if campi_extra:
            payload_completo.update(campi_extra)

        # caso ideale: campi_canonici presente
        if campi_canonici:
            payload_completo.update(campi_canonici)

        # caso fallback 1: wrapper "data"
        if not payload_completo and isinstance(data.get("data"), dict):
            payload_completo.update(data["data"])

        # caso fallback 2: root unico con dict dentro
        if not payload_completo and len(data) == 1:
            only_value = next(iter(data.values()))
            if isinstance(only_value, dict):
                payload_completo.update(only_value)

        # caso fallback 3: prendi tutto il dict, esclusi meta-campi tecnici
        if not payload_completo:
            payload_completo = {
                k: v for k, v in data.items()
                if k not in {"tipo_documento", "campi_canonici", "campi_extra"}
            }

        print("[GPT][SCHEMA] Payload completo estratto:", payload_completo)

        return tipo_documento, payload_completo

    except Exception as e:
        print(f"Errore GPT schema parsing: {e}")
        return document_type, {}
# --------------------------
# Generazione Word
# --------------------------
def generate_word(fields, original_file):
    os.makedirs("output", exist_ok=True)
    doc = Document()
    doc.add_heading("Dati Documento", 0)
    for campo, valore in fields.items():
        if isinstance(valore,(dict,list)):
            valore = json.dumps(valore, ensure_ascii=False)
        doc.add_paragraph(f"{campo}: {valore}")
    output_path = os.path.join("output", os.path.basename(original_file)+".docx")
    doc.save(output_path)
    print(f"[WORD] Documento Word generato: {output_path}")
    return output_path

# --------------------------
# LEGACY / NON USATO DAL FLUSSO ATTUALE
# Queste funzioni appartengono al vecchio approccio di ricerca/parsing.
# Il flusso attuale usa:
# - parse_document_with_gpt(...)
# - parse_search_query_with_gpt(...) in query_understanding_gpt.py
# - search_documents() in api_server.py
# --------------------------


# --------------------------
# AI Search nei documenti
# --------------------------
def interpret_search_query(query):
    """
    Interpreta la query dell'utente e restituisce filtri
    per la ricerca nei documenti.
    """
    system_prompt = """
Sei un motore di ricerca per documenti.

Interpreta la richiesta dell'utente e restituisci i filtri
per cercare nel database.

Rispondi SOLO con JSON nel formato:

{
  "tipo_documento": "stringa oppure null",
  "campo": "nome campo oppure null",
  "valore": "valore oppure null",
  "operatore": ">, <, >=, <=, = oppure null"
}

Regole:
- "tipo_documento" va valorizzato solo se la query parla chiaramente di un tipo documento.
- "campo" è il nome del dato da filtrare.
- "valore" è il valore cercato.
- "operatore" va valorizzato solo se la query contiene un confronto numerico o temporale.
- Se non c'è confronto esplicito, usa null.
- Se l'utente cerca per testo semplice, metti operatore = "=" oppure null.
- Restituisci SEMPRE JSON valido.
- Non aggiungere spiegazioni.

Esempi:

Query: passaporti italiani
Risposta:
{
  "tipo_documento": "passaporto",
  "campo": "nazionalita",
  "valore": "Italia",
  "operatore": null
}

Query: nome Mario
Risposta:
{
  "tipo_documento": null,
  "campo": "nome",
  "valore": "Mario",
  "operatore": null
}

Query: biglietti traghetto
Risposta:
{
  "tipo_documento": "biglietto traghetto",
  "campo": null,
  "valore": null,
  "operatore": null
}

Query: retribuzione maggiore di 40000
Risposta:
{
  "tipo_documento": null,
  "campo": "retribuzione",
  "valore": "40000",
  "operatore": ">"
}

Query: documenti con ore settimanali >= 40
Risposta:
{
  "tipo_documento": null,
  "campo": "ore_settimanali",
  "valore": "40",
  "operatore": ">="
}

Query: contratti con data dopo 2024-01-01
Risposta:
{
  "tipo_documento": "contratto",
  "campo": "data",
  "valore": "2024-01-01",
  "operatore": ">"
}
"""
    try:
        response = openai.chat.completions.create(
            model="gpt-4o-mini",
            temperature=0,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": query}
            ]
        )

        content = response.choices[0].message.content.strip()

        if content.startswith("```"):
            content = content.replace("```json", "").replace("```", "").strip()

        data = json.loads(content)

        return {
            "tipo_documento": data.get("tipo_documento"),
            "campo": data.get("campo"),
            "valore": data.get("valore"),
            "operatore": data.get("operatore")
        }

    except Exception as e:
        print("Errore AI search:", e)
        return {
            "tipo_documento": None,
            "campo": None,
            "valore": None,
            "operatore": None
        }
# --------------------------
# Filtra documenti lato server usando query AI
# --------------------------
def filter_documents_by_query(docs, query):
    """
    Applica la ricerca AI sui documenti già presenti nel DB.
    """
    filters = interpret_search_query(query)
    tipo_doc_filter = filters.get("tipo_documento")
    campo_filter = filters.get("campo")
    valore_filter = filters.get("valore")

    filtered_docs = []
    for d in docs:
        if tipo_doc_filter and d["tipo_documento"] != tipo_doc_filter:
            continue
        if campo_filter and valore_filter:
            field_val = d["campi"].get(campo_filter)
            if not field_val:
                continue
            if valore_filter.lower() not in str(field_val).lower():
                continue
        filtered_docs.append(d)
    return filtered_docs

