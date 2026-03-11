import os
import time
import shutil
from datetime import datetime
from watchdog.observers import Observer
from watchdog.events import FileSystemEventHandler
from PIL import Image
import pytesseract
from docx import Document
from pdf2image import convert_from_path
import openai

# --------------------------
# CONFIGURAZIONE API GPT
# --------------------------
openai.api_key = os.getenv("OPENAI_API_KEY")  # Legge la chiave dall'ambiente

# --------------------------
# PERCORSO TESSERACT
# --------------------------
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

# --------------------------
# DATABASE
# --------------------------
from db_setup import engine, documents, document_fields

# --------------------------
# CARTELLE
# --------------------------
INPUT_FOLDER = "input/"
PROCESSED_FOLDER = "processed/"
OUTPUT_FOLDER = "output/"

# ==========================
# FUNZIONI DATABASE
# ==========================
def save_document(tipo_documento, file_path, fields):
    with engine.begin() as conn:
        result = conn.execute(documents.insert().values(
            tipo_documento=tipo_documento,
            file_path=file_path,
            data_creazione=datetime.now()
        ))
        document_id = result.inserted_primary_key[0]

        for campo, valore in fields.items():
            conn.execute(document_fields.insert().values(
                document_id=document_id,
                campo=campo,
                valore=valore
            ))
    return document_id

# ==========================
# FUNZIONI OCR
# ==========================
def extract_text(file_path):
    try:
        text = pytesseract.image_to_string(Image.open(file_path), lang="ita+eng")
        return text
    except Exception as e:
        print(f"Errore OCR su {file_path}: {e}")
        return ""

def extract_text_from_pdf(pdf_path):
    try:
        text_all = ""
        pages = convert_from_path(pdf_path, poppler_path=r"C:\Program Files\poppler-25.12.0\bin")
        for page in pages:
            text_all += pytesseract.image_to_string(page, lang="ita+eng") + "\n"
        return text_all
    except Exception as e:
        print(f"Errore OCR su PDF {pdf_path}: {e}")
        return ""

# ==========================
# FUNZIONE GPT PER PARSING INTELLIGENTE
# ==========================
def parse_data_with_gpt(text, tipo_documento="generico"):
    """
    Usa GPT per trasformare il testo OCR in un dizionario chiave-valore
    (nuova API openai>=1.0)
    """
    try:
        prompt = f"""
        Estrai tutti i dati significativi dal seguente testo di documento {tipo_documento}.
        Organizza i dati in un dizionario JSON chiave-valore.
        Esempio:
        {{
            "nome": "...",
            "cognome": "...",
            "data_nascita": "...",
            "codice_fiscale": "...",
            "indirizzo": "...",
            "telefono": "...",
            "email": "...",
            "iban": "...",
            "importo": "..."
        }}
        Testo:
        \"\"\"{text}\"\"\"
        """

        response = openai.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "user", "content": prompt}],
            temperature=0
        )

        content = response.choices[0].message.content
        import json
        fields = json.loads(content)
        return fields
    except Exception as e:
        print(f"Errore GPT parsing: {e}")
        return {}

# ==========================
# GENERAZIONE WORD
# ==========================
def generate_word(fields, original_file):
    doc = Document()
    doc.add_heading("Dati Documento", 0)
    for campo, valore in fields.items():
        doc.add_paragraph(f"{campo}: {valore}")
    output_path = os.path.join(OUTPUT_FOLDER, os.path.basename(original_file) + ".docx")
    doc.save(output_path)
    print(f"Documento Word generato: {output_path}")

# ==========================
# PROCESSO COMPLETO
# ==========================
def process_document(file_path):
    time.sleep(1)
    print(f"Elaborazione file: {file_path}")

    text = ""
    if file_path.lower().endswith((".jpg", ".jpeg", ".png")):
        text = extract_text(file_path)
    elif file_path.lower().endswith(".pdf"):
        text = extract_text_from_pdf(file_path)
    else:
        print(f"Formato file non supportato: {file_path}")
        return

    # Parsing intelligente con GPT
    fields = parse_data_with_gpt(text)

    # Salva i dati nel database
    save_document("generico", file_path, fields)

    # Genera documento Word
    generate_word(fields, file_path)

    # Sposta file originale
    shutil.move(file_path, os.path.join(PROCESSED_FOLDER, os.path.basename(file_path)))
    print(f"File spostato in processed/")

# ==========================
# WATCHER CARTELLA
# ==========================
class Watcher(FileSystemEventHandler):
    def on_created(self, event):
        if not event.is_directory:
            process_document(event.src_path)

# ==========================
# MAIN
# ==========================
if __name__ == "__main__":
    os.makedirs(INPUT_FOLDER, exist_ok=True)
    os.makedirs(PROCESSED_FOLDER, exist_ok=True)
    os.makedirs(OUTPUT_FOLDER, exist_ok=True)

    observer = Observer()
    observer.schedule(Watcher(), INPUT_FOLDER, recursive=False)
    observer.start()
    print("Watcher avviato... metti un file in input/ per test")
    try:
        while True:
            time.sleep(1)
    except KeyboardInterrupt:
        observer.stop()
    observer.join()